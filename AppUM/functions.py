from django.shortcuts import render
from .models import *
from .utils import crear_pdf
from django.core.paginator import Paginator
from django.shortcuts import render
from .models import *
from django.core.files.base import ContentFile
from datetime import datetime
from.inserts import *
from docx import Document
import locale
import win32com.client
import pythoncom
from django.conf import settings
import os
from django.core.files import File
import json
from django.core.mail import send_mail
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from django.http import HttpResponse
from num2words import num2words
import time
import shutil
import threading


def departamento_matricula(matricula):
    primer_letra = matricula[0:1:1]
    if primer_letra == "A":
        departamento = "Canelones"
    elif primer_letra == "B":
        departamento = "Maldonado"
    elif primer_letra == "C":
        departamento = "Rocha"
    elif primer_letra == "D":
        departamento = "Treina y Tres"
    elif primer_letra == "E":
        departamento = "Cerro Largo"
    elif primer_letra == "F":
        departamento = "Rivera"
    elif primer_letra == "G":
        departamento = "Artigas"
    elif primer_letra == "H":
        departamento = "Salto"
    elif primer_letra == "I":
        departamento = "Paysandú"
    elif primer_letra == "J":
        departamento = "Río Negro"
    elif primer_letra == "K":
        departamento = "Soriano"
    elif primer_letra == "L":
        departamento = "Colonia"
    elif primer_letra == "M":
        departamento = "San José"
    elif primer_letra == "N":
        departamento = "Flores"
    elif primer_letra == "O":
        departamento = "Florida"
    elif primer_letra == "P":
        departamento = "Lavalleja"
    elif primer_letra == "Q":
        departamento = "Durazno"
    elif primer_letra == "R":
        departamento = "Tacuarembó"
    elif primer_letra == "S":
        departamento = "Montevideo"
    else: 
        departamento = None
    return departamento

def pdf_crear(req,ruta_pdf,renderizar_en,model_insert,datos_a_pdf,nombre_archivo,mensaje,negocio):
    if negocio == "UM":
        logo = Logos.objects.get(id=1)
        logo_um_url = req.build_absolute_uri(logo.logo_UM.url) if logo.logo_UM else None

        logo_contexto = logo.logo_UM.url if logo.logo_UM else None
    else:
        logo = Logos.objects.get(id=2)
        logo_dm_url = req.build_absolute_uri(logo.logo_DM.url) if logo.logo_DM else None
        logo_contexto = logo.logo_DM.url if logo.logo_DM else None
        
    pdf = crear_pdf(ruta_pdf, datos_a_pdf)             
    if pdf:
         pdf_file = ContentFile(pdf.content)
         file_name = nombre_archivo
 # Asignar el archivo al campo identificacion_pdf de la moto
         model_insert.identificacion_pdf.save(file_name, pdf_file)
         model_insert.save()

    contexto = {
                "messages":mensaje,
                'pdf_url': model_insert.identificacion_pdf.url,
                "logo_um": logo_contexto

                }
    return render(req,"perfil_administrativo/motos/contenido_pdf.html",contexto)

def contexto_para_pdf_moto(model_moto,logo):
        datos_a_pdf = {
                        "moto": {
                            "id":model_moto.id,
                            "marca": model_moto.marca,
                            "modelo": model_moto.modelo,
                            "motor": model_moto.motor,
                            "anio": model_moto.anio,
                            "kilometros": model_moto.kilometros,
                            "precio": model_moto.precio
                        },
                        "current_year": datetime.now().year,
                            "logo_um": logo  # Asegúrate de importar datetime
                        }
        
        return datos_a_pdf

def checkbox_pdf(req,nueva_moto):
    ruta_pdf = "perfil_administrativo/motos/identificacion_moto.html"
    logo_um = Logos.objects.get(id=1)
    logo_um_url = req.build_absolute_uri(logo_um.logo_UM.url) if logo_um.logo_UM else None
    datos_a_pdf = contexto_para_pdf_moto(nueva_moto,logo_um_url)
    renderizar_en = "perfil_administrativo/motos/contenido_pdf.html"
    nombre_archivo = f"identificacion_{nueva_moto.id}.pdf"
    mensaje = "Moto ingresada con éxito. A continuación se visualiza el PDF generado para identificar fisicamente la misma."
    pdf_ret = pdf_crear(req,ruta_pdf,renderizar_en,nueva_moto,datos_a_pdf,nombre_archivo,mensaje,"UM")
    return pdf_ret

def validacion_moto(id_moto,num_motor,num_chasis):
    moto_actual = Moto.objects.filter(id = id_moto).first()
    existe_num_motor = Moto.objects.filter(num_motor = num_motor).first()
    existe_num_chasis = Moto.objects.filter(num_chasis = num_chasis).first()
    
    if existe_num_motor:
        if existe_num_motor.num_motor != moto_actual.num_motor:
            return "existe_num_motor"
    
    if existe_num_chasis:
        if existe_num_chasis.num_chasis != moto_actual.num_chasis:
            return "existe_num_chasis"
        
def matricula_valid(matricula,padron,id_moto):
    existe_matricula = Matriculas.objects.filter(matricula = matricula).first()
    # moto = Moto.objects.filter(num_motor = num_motor,num_chasis = num_chasis).first()
    if existe_matricula:
        if existe_matricula.moto_id != id_moto:
            return "matricula_existe"
    
    existe_padron = Matriculas.objects.filter(padron=padron).first()
    if existe_padron:
        if existe_padron.moto_id != id_moto:
            return "padron_existe"

# def valid_cliente(documento,tel1,tel2,correo1,correo2):
def valid_cliente(documento):
    existe_cliente = Cliente.objects.filter(documento=documento).first()
    if existe_cliente:
        return "existe_cliente"
    
    # existe_tel_1 = ClienteTelefono.objects.filter(telefono=tel1).first()

    # if existe_tel_1:
    #     return "existe_telefono_1"
    
    # if tel2:
    #     existe_tel_2 = ClienteTelefono.objects.filter(telefono=tel2).first()
    #     if existe_tel_2:
    #         return "existe_telefono_2"
    

    # if correo1:
    #     existe_correo_1 = ClienteCorreo.objects.filter(correo=correo1).first()
    #     if existe_correo_1:
    #         return "existe_correo_1"
        
    # if correo2:
    #     existe_correo_2 = ClienteCorreo.objects.filter(correo=correo2).first()
    #     if existe_correo_2:
    #         return "existe_correo_2"

def contexto_para_cliente(id_cliente,mensaje_error):
            cliente = Cliente.objects.get(id=id_cliente)
            tel_princ = ClienteTelefono.objects.get(principal=1,cliente_id=id_cliente)
            tel_sec = ClienteTelefono.objects.filter(principal=0,cliente_id=id_cliente).first()

            tipo_documento_ci = cliente.documento[0:2:1]
            tipo_documento_pas_dni = cliente.documento[0:3:1]

            longitud_doc = len(cliente.documento)
            doc_num = ""
            
            if tipo_documento_ci == "CI":
                #
                tipo_doc = "CI"
                for i in range(2,longitud_doc):
                    doc_num = doc_num + cliente.documento[i]
            elif tipo_documento_pas_dni == "DNI":
                tipo_doc = "DNI"
                for i in range(3,longitud_doc):
                    doc_num = doc_num + cliente.documento[i]
            else:
                tipo_doc = "PAS"
                for i in range(3,longitud_doc):
                    doc_num = doc_num + cliente.documento[i]
            
            correo_princ = ClienteCorreo.objects.filter(principal=1,cliente_id=id_cliente).first()
            correo_sec = ClienteCorreo.objects.filter(principal=0,cliente_id=id_cliente).first()

            if tel_sec:
                t_s = tel_sec.telefono
            else:
                t_s = None

            if correo_princ:
                longitud_correo_princ = len(correo_princ.correo)
                c_p = "" 
                lon = 0
                for i in range(0,longitud_correo_princ):
                    if correo_princ.correo[i] == "@":
                        j = i
                        break
                    else:
                        lon += 1
                        c_p = c_p + correo_princ.correo[i]
                
                longitud_dominio = longitud_correo_princ - lon
                dom_princ = ""
                pos = j
                for j in range(pos,longitud_correo_princ):
                    dom_princ = dom_princ + correo_princ.correo[j]
            else:
                c_p = ""
                dom_princ = None
            
            if correo_sec:
                
                longitud_correo_sec = len(correo_sec.correo)
                c_s = "" 
                lon = 0
                for i in range(0,longitud_correo_sec):
                    if correo_sec.correo[i] == "@":
                        j = i
                        break
                    else:
                        lon += 1
                        c_s = c_s + correo_sec.correo[i]
                
            
                dom_sec = ""
                pos = j
                for j in range(pos,longitud_correo_sec):
                    dom_sec = dom_sec + correo_sec.correo[j]
            else:
                c_s = ""
                dom_sec = None
            f_nac_formateada = cliente.fecha_nacimiento.strftime('%Y-%m-%d') if cliente.fecha_nacimiento else None
            contexto = {"datos_cliente":cliente,
                        "tipo_doc":tipo_doc,
                        "doc_num":doc_num,
                        "tel_princ":tel_princ,
                        "fecha_nac":f_nac_formateada,
                        "tel_sec":t_s,
                        "correo_princ":c_p,
                        "dom_princ":dom_princ,
                        "dom_sec":dom_sec,
                        "correo_sec":c_s,
                        "error_message":mensaje_error}
            return contexto

# def valid_cliente_mod(id_cliente,documento,tel1,tel2,correo1,correo2):
def valid_cliente_mod(id_cliente,documento):
    cliente_id = Cliente.objects.get(id=id_cliente)
    cliente_doc = Cliente.objects.filter(documento=documento).first()

    if cliente_doc:
        if id_cliente != cliente_doc.id:
            return "existe_cliente"
    
    # tel_principal = ClienteTelefono.objects.filter(telefono=tel1).first()
    # if tel_principal:
    #     if tel_principal.cliente_id != id_cliente:
    #         return "existe_tel_principal"
    
    # tel_secundario = ClienteTelefono.objects.filter(telefono=tel2).first()
    # if tel_secundario:
    #     if tel_secundario.cliente_id != id_cliente:
    #         return "existe_tel_secundario"
        
    # correo_princ = ClienteCorreo.objects.filter(correo=correo1).first()
    # if correo_princ:
    #     if correo_princ.cliente_id != id_cliente:
    #         return "existe_correo_principal"

    # correo_sec = ClienteCorreo.objects.filter(correo=correo2).first()
    # if correo_sec:
    #     if correo_sec.cliente_id != id_cliente:
    #         return "existe_correo_secundario"

def validar_personal(documento,telefono,correo):
    personal = Personal.objects.filter(documento=documento).first()
    if personal:
        #VERIFICAR SI EXISTE EN ADMINISTRATIVO
        administrativo = Administrativo.objects.filter(personal_ptr_id = personal.id).first()
        if administrativo:
            if administrativo.activo == 1: #SI EXISTIERA DEVUELVE ERROR 
                return "existe_admin"
            else: #SI EXISTIERA Y FUE DADO DE BAJA COMO ADMINISTRATIVO (EJ. FUE ADMIN, PASO A SER MECANICO Y VUELVE A SER ADMIN), SIMPLEMENTE SE MODIFICA EL CAMPO ACTIVO = 1
                return "admin_desactivado"
        
        mecanico = Mecanico.objects.filter(personal_ptr_id = personal.id).first()
        if mecanico:
            if mecanico.activo == 1: #SI EXISTIERA DEVUELVE ERROR 
                return "existe_mecanico"
            else: #SI EXISTIERA Y FUE DADO DE BAJA COMO ADMINISTRATIVO (EJ. FUE ADMIN, PASO A SER MECANICO Y VUELVE A SER ADMIN), SIMPLEMENTE SE MODIFICA EL CAMPO ACTIVO = 1
                return "mecanico_desactivado"
    
    telefono_pers = Personal.objects.filter(telefono = telefono).first()

    if telefono_pers: #SI EL TELEFONO DEL PERSONAL A INGRESAR EXISTIERA (PUEDE QUE SEA UN MECANICO QUE SE LE ASIGNARA TAREAS DE ADMINISTRATIVO), EL ADMINISTRATIVO LOGEADO EN EL SISTEMA DEBERA OTORGARLE PERMISOS DE ADMINISTRATIVO
        # tel_admin = Administrativo.objects.filter(ptr_personal_id = telefono_pers.id).first()
        # if tel_admin:
        #     if tel_admin.activo == 1: 
                return "existe_telefono" 
    
    correo_pers = Personal.objects.filter(correo = correo).first()
    if correo_pers: #SI EL CORREO DEL PERSONAL A INGRESAR EXISTIERA (PUEDE QUE SEA UN MECANICO QUE SE LE ASIGNARA TAREAS DE ADMINISTRATIVO), EL ADMINISTRATIVO LOGEADO EN EL SISTEMA DEBERA OTORGARLE PERMISOS DE ADMINISTRATIVO
        # correo_admin = Administrativo.objects.filter(ptr_personal_id = correo_pers.id).first()
        # if correo_admin:
        #     if correo_admin.activo == 1:
                return "existe_correo"

def nombre_usuario(nombre,apellido):
    #SI EXISTE EL USUARIO, AGREGAR NUMERO AL FINAL (EJ. jperez, jperez1, jperez2)
    apellido_unido = apellido.replace(' ','')
    usuario = nombre[0:1:1] + apellido_unido
    cant_usuario = Personal.objects.filter(usuario__contains=usuario).count()
    if cant_usuario > 0: #EXISTE MAS DE UN jperez
        usuario = nombre[0:1:1] + apellido + str(cant_usuario)
    usuario_minuscula = usuario.lower()
    return usuario_minuscula

def validar_billetes(total_billetes_2000,total_billetes_1000,total_billetes_500,total_billetes_200,total_billetes_100,total_billetes_50,total_billetes_20):
    if total_billetes_2000 < 0:
        error = "Ingrese un valor correcto para los billetes de $2000"
    elif total_billetes_1000 < 0:
        error = "Ingrese un valor correcto para los billetes de $1000"
    elif total_billetes_500 < 0:
        error = "Ingrese un valor correcto para los billetes de $500"
    elif total_billetes_200 < 0:
        error = "Ingrese un valor correcto para los billetes de $200"
    elif total_billetes_100 < 0:
        error = "Ingrese un valor correcto para los billetes de $100"
    elif total_billetes_50 < 0:
        error = "Ingrese un valor correcto para los billetes de $50"
    elif total_billetes_20 < 0:
        error = "Ingrese un valor correcto para los billetes de $20"
    else:
        error = None
    return error

def validar_monedas(total_monedas_50,total_monedas_10,total_monedas_5,total_monedas_2,total_monedas_1):
    if total_monedas_50 < 0:
        error = "Ingrese un valor correcto para las monedas de $50"
    elif total_monedas_10 < 0:
        error = "Ingrese un valor correcto para las monedas de $10"
    elif total_monedas_5 < 0:
        error = "Ingrese un valor correcto para las monedas de $5"
    elif total_monedas_2 < 0:
        error = "Ingrese un valor correcto para las monedas de $2"
    elif total_monedas_1 < 0:
        error = "Ingrese un valor correcto para las monedas de $1"
    else:
        error = None
    return error

def funcion_paginas_varias(req,instancia_model):
    paginator = Paginator(instancia_model, 5)  # 5 motos por página

    page_number = req.GET.get('page')  # Obtiene el número de página desde la URL
    page_obj = paginator.get_page(page_number)

    return page_obj

def validar_entrega_menor_precio(moneda_entrega,entrega,elemento,precio_dolar,elemento_tipo,existe_cuota,id_elemento):
    error = None
    if moneda_entrega == "Pesos":             
        if elemento.moneda_precio == "Pesos":
            if int(entrega) > int(elemento.precio):
                error = "El pago no puede exceder el precio total del accesorio"
        else:
            if int(entrega) > int((elemento.precio * precio_dolar)):
                error = "El pago no puede exceder el precio total del accesorio"            
    else:
        if elemento.moneda_precio == "Pesos":
            if int(int(entrega) * precio_dolar) > int(elemento.precio):
                error = "El pago no puede exceder el precio total del accesorio"
        else:
            if int(entrega) > int((elemento.precio * precio_dolar)):
                error = "El pago no puede exceder el precio total del accesorio" 
    
    if existe_cuota:
        if elemento_tipo == "Accesorio":
            cuota = CuotasAccesorios.objects.filter(venta_id=id_elemento).latest('id')
        else:
            cuota = CuotasMoto.objects.filter(venta_id=id_elemento).latest('id')
        
        if moneda_entrega == "Pesos":
            if int(entrega) > int(cuota.cant_restante_pesos):
                error = "El pago no puede exceder el valor restante en pesos"
        else:
            if int(entrega) > int(cuota.cant_restante_dolares):
                error = "El pago no puede exceder el valor restante en dólares"
    
    return error

def obtener_compras_accesorios(req,codigo_compra):
        venta = ClienteAccesorio.objects.filter(codigo_compra=codigo_compra).latest('id')
        id_venta = venta.id
        forma_pago = venta.forma_de_pago
        #
        resultados_cuotas = (
                CuotasAccesorios.objects
                .filter(venta_id=id_venta)
                .values(
                    'id',
                    'fecha_pago', 
                     'metodo_pago', 
                    'cant_restante_dolares', 
                    'cant_restante_pesos', 
                    'moneda', 
                    'observaciones',
                    'valor_pago_dolares',
                    'valor_pago_pesos',
                    'comprobante_pago'
                )
            )
            #
        res_pagos = []
        i = 1
        for resultado in resultados_cuotas:
                    ca = CuotasAccesorios.objects.get(id=resultado['id'])
                    res_pagos.append({
                    'cuota': resultado,
                    'comprobante_pago': ca.comprobante_pago.url if ca.comprobante_pago else None,
                    'mostrar_boton': i == len(resultados_cuotas)           
                    })
                    i = i + 1
        venta = ClienteAccesorio.objects.get(id=id_venta)
        cliente = Cliente.objects.get(id=venta.cliente_id)
        telefono = ClienteTelefono.objects.filter(cliente=cliente,principal=1).first()
        cliente_data = {
        "cliente": cliente.nombre + " " + cliente.apellido,
        "telefono": telefono.telefono,
        "direccion": cliente.domicilio,
        }
        cliente_json = json.dumps(cliente_data)

        venta_accesorios = ClienteAccesorio.objects.filter(codigo_compra=codigo_compra)
        dolar = PrecioDolar.objects.get(id=1)
        precio_dolar = float(dolar.precio_dolar_tienda)
        accesorio_data = []
        total_pesos = 0
        total_dolares = 0
        for a in venta_accesorios:
            if int(a.codigo_compra) >= 1:
                accesorio = Accesorio.objects.get(id=a.accesorio_id)
                # talle = "Talle " + accesorio.talle if accesorio.talle != "Sin talle" else None
                moneda = "$" if accesorio.moneda_precio == "Pesos" else "U$S"
                if accesorio.moneda_precio == "Pesos":
                    precio_en_pesos = int(accesorio.precio)
                    precio_en_dolares = int(accesorio.precio) / precio_dolar
                else:
                    precio_en_dolares = int(accesorio.precio)
                    precio_en_pesos = int(precio_en_dolares) * precio_dolar
                total_pesos = total_pesos + precio_en_pesos
                total_dolares = total_dolares + precio_en_dolares

                if accesorio.talle == "Sin talle":
                    detalle = accesorio.tipo + " " + accesorio.marca + " " + accesorio.modelo
                else:
                    detalle = accesorio.tipo + " " + accesorio.marca + " " + accesorio.modelo + " Talle " + accesorio.talle
                
                accesorio_data.append({
                    "detalle":detalle,
                    "precio": moneda + str(accesorio.precio),
                    "id_accesorio":accesorio.id
                })
                
                accesorios_para_json = [
                    {
                            "detalle":p["detalle"],
                            "precio": p["precio"]
                    }
                    for p in accesorio_data
                ]

        accesorio_json = json.dumps(accesorios_para_json)
        
        pagos = CuotasAccesorios.objects.filter(venta_id=id_venta,mostrar_reporte=1).order_by('-id')
        if pagos:
            p_pagos_data = [
                    {   
                        "fecha":cuota.fecha_pago.strftime('%Y-%m-%d'),
                        "moneda":cuota.moneda,
                        "monto": float(cuota.valor_pago_pesos) if cuota.moneda == "Pesos" else float(cuota.valor_pago_dolares),
                        "metodo":cuota.metodo_pago
                    }
                    for cuota in pagos
                ]
            pagos_accesorios_json = json.dumps(p_pagos_data)
        else:
            pagos_accesorios_json = None

        page_obj = funcion_paginas_varias(req,res_pagos)
        # total_pesos_json = json.dumps(total_pesos)
        # total_dolares_json = json.dumps(total_dolares)
        data_total_precios = {
            "total_pesos": round(total_pesos, 2),
            "total_dolares": round(total_dolares, 2)
        }
        total_precios_json = json.dumps(data_total_precios)
        data = [
            accesorio_data,
            total_pesos,
            total_dolares,
            id_venta,
            page_obj,
            cliente_json,
            accesorio_json,
            pagos_accesorios_json,
            total_precios_json,
            venta.cliente_id,
            forma_pago
        ]

        return data

def obtener_compras_motos(req,id_cv):
        resultados_cuotas = (
                CuotasMoto.objects
                .filter(venta_id=id_cv)
                .values(
                    'id',
                    'fecha_pago', 
                    'fecha_prox_pago',  
                    'cant_restante_dolares', 
                    'cant_restante_pesos', 
                    'moneda', 
                    'observaciones',
                    'valor_pago_dolares',
                    'valor_pago_pesos',
                    'comprobante_pago'
                )
            )
            
        res_documentacion = []
        i = 1
        for resultado in resultados_cuotas:
                    cm = CuotasMoto.objects.get(id=resultado['id'])
                    res_documentacion.append({
                    'cuota': resultado,
                    'comprobante_pago': cm.comprobante_pago.url if cm.comprobante_pago else None,
                    'mostrar_boton': i == len(resultados_cuotas)           
                    })
                    i = i + 1
        page_obj = funcion_paginas_varias(req,res_documentacion)
        return page_obj

def valores_compras(existe_cuota,moneda,entrega,id_elemento,elemento,elemento_tipo,precio_dolar):
    if not existe_cuota:
        if moneda == "Pesos":
            entrega_pesos = entrega
            entrega_dolares = 0
            if elemento.moneda_precio == "Pesos":
                resto_pesos = int(elemento.precio) - int(entrega_pesos)
                resto_dolares = resto_pesos / precio_dolar
            else:
                resto_pesos = int((elemento.precio * precio_dolar)) - int(entrega_pesos)
                resto_dolares = resto_pesos / precio_dolar                       
        else:
            entrega_pesos = 0
            entrega_dolares = entrega
            if elemento.moneda_precio == "Pesos":
                resto_dolares = int((elemento.precio / precio_dolar)) - int(entrega_dolares)
                resto_pesos = resto_dolares * precio_dolar
            else:
                resto_dolares = int(elemento.precio) - int(entrega_dolares)
                resto_pesos = resto_dolares * precio_dolar 
    else:
        if elemento_tipo == "Accesorio":
            cuota = CuotasAccesorios.objects.filter(venta_id=id_elemento).latest('id')
        else:
            cuota = CuotasMoto.objects.filter(venta_id=id_elemento).latest('id')
        
        if moneda == "Pesos":
            entrega_pesos = entrega
            entrega_dolares = 0
            resto_pesos = int(cuota.cant_restante_pesos) - int(entrega_pesos)
            resto_dolares = resto_pesos / precio_dolar
        else:
            entrega_pesos = 0
            entrega_dolares = entrega
            resto_dolares = int(cuota.cant_restante_dolares) - int(entrega_dolares)
            resto_pesos = resto_dolares * precio_dolar

            #resto_dolares,resto_pesos,entrega_pesos,entrega_dolares
    
    lista = [resto_dolares,resto_pesos,entrega_pesos,entrega_dolares]
    return lista

def valores_pagos_accesorios():
    pass

def obtener_detalles_cuotas_comunes(id_cv):
    compra = ComprasVentas.objects.get(id=id_cv)
    moto = Moto.objects.get(id=compra.moto_id)
    producto = f"{moto.marca} {moto.modelo}"
    primeros_pagos = CuotasMoto.objects.filter(venta_id=id_cv,tipo_pago__in=["Seña", "Entrega inicial","Entrega"]).order_by('-id')
    existen_pagos = CuotasMoto.objects.filter(venta_id=id_cv).first()
    if existen_pagos:
        ultimo_pago_general = CuotasMoto.objects.filter(venta_id=id_cv).latest('id')
        data_p_pagos = []
        for p_p in primeros_pagos:
            data_p_pagos.append({
                "primeros_pagos":p_p,
                "boton_ultimo_pago": True if p_p.id == ultimo_pago_general.id else False
            })
        # moneda = "$" if moto.moneda_precio == "Pesos" else "U$S"
        # primer_fin = Financiamientos.objects.filter(venta_id=id_cv).order_by('id').first()
        # financiamiento = f"{str(primer_fin.cantidad_cuotas)} x {moneda} {str(primer_fin.valor_cuota)}"
    else:
        data_p_pagos = None
        financiamiento = None


    #LISTAR FINANCIAMIENTOS EN SELECT
    financiamientos_select = Financiamientos.objects.filter(venta_id=id_cv,inicial=0).order_by('-id')
    data_financiamientos = []
    for f_s in financiamientos_select:
        fecha_str = f_s.fecha.strftime("%d/%m/%Y")  # Ajusta el formato si es necesario
        actual = "Actual" if f_s.actual else ""
        data_financiamientos.append({
            "financiamientos": f"{fecha_str} {actual}".strip(),  # Quita espacios en blanco si no hay "Actual"
            "id": f_s.id,
        })
    
    data = [
        producto,
        moto.precio,
        moto.precio_final,
        compra.forma_de_pago,
        data_p_pagos,
        # financiamiento,
        data_financiamientos,
    ]
    return data

def obtener_detalles_cuotas_financiamiento(req,id_f):
    resultados = (
                CuotasFinanciacion.objects
                .filter(financiamiento_id=id_f)
                .select_related('cuota','financiamiento')
                .values(
                    'cuota__id',
                    'cuota__fecha_pago', 
                    'cuota__fecha_prox_pago',  
                    'cuota__cant_restante_dolares', 
                    'cuota__cant_restante_pesos', 
                    'cuota__moneda', 
                    'cuota__observaciones',
                    'cuota__valor_pago_dolares',
                    'cuota__valor_pago_pesos',
                    'cuota__comprobante_pago',
                    'cuota__tipo_pago',
                    'cuota__metodo_pago',
                    'financiamiento__id'
                ).order_by('-cuota__id')
            )
    res_pagos = []
    f = Financiamientos.objects.get(id=id_f)
    existen_pagos = CuotasMoto.objects.filter(venta_id=f.venta_id).first()
    if existen_pagos:
        ult_cuota = CuotasMoto.objects.filter(venta_id=f.venta_id).latest('id')
        i = 1
        for resultado in resultados:
            ca = CuotasMoto.objects.get(id=resultado['cuota__id'])
            f = Financiamientos.objects.get(id=resultado['financiamiento__id'])
            res_pagos.append({
            'cuota': resultado,
            'comprobante_pago': ca.comprobante_pago.url if ca.comprobante_pago else None,
            # 'mostrar_boton': (i == len(resultados) and f.actual) or (ca.id == ult_cuota.id)  
            'mostrar_boton': (ca.id == ult_cuota.id) and f.actual
            })
            i = i + 1

    page_obj = funcion_paginas_varias(req,res_pagos)
    return page_obj

def crear_certificado_bikeup(cliente,telefono,moto,id_cv):
    output_dir = os.path.join(settings.MEDIA_ROOT, 'certificado_venta')

    # doc.save(docx_file_path)
    # doc = Document('D:\Escritorio\SISTEMA UMPIERREZ MOTOS\cert_bikeup.docx')  # Este es el archivo de Word que quieres editar
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Ruta del archivo de plantilla (ajustado para que sea desde media)
    plantilla_path = os.path.join(settings.BASE_DIR, 'media', 'cert_bikeup.docx')
    
    # Crear el documento Word a partir del archivo de plantilla
    doc = Document(plantilla_path)

    locale.setlocale(locale.LC_TIME, 'spanish')
    fecha_actual = datetime.now()
    fecha_formateada = fecha_actual.strftime("%d DE %B").upper()
    longitud_doc = len(cliente.documento)
    doc_num = ""
    # apto = "APTO " + str(cliente.num_apartamento) if cliente.num_apartamento > 0 else ""
    # direccion = cliente.calle.upper() + " " + str(cliente.numero) + apto + ", " + cliente.ciudad.upper()
    for i in range(2,longitud_doc):
                    doc_num = doc_num + cliente.documento[i]
    if moto.moneda_precio == "Pesos":
        moneda = "PESOS"
    else:
        moneda = "DOLARES"
    
    domicilio = cliente.domicilio.upper()
    if moto.contiene_num_motor:
        numero_motor = moto.num_motor
    else:
        numero_motor = "SIN NUMERO DE MOTOR"

    if moto.contiene_num_chasis:
        numero_chasis = moto.num_chasis
    else:
        numero_chasis = "SIN NUMERO DE CHASIS"
    for p in doc.paragraphs:
        if 'FECHA_ACTUAL' in p.text:
            p.text = p.text.replace('FECHA_ACTUAL', fecha_formateada)
        if 'NOMBRE_CLIENTE' in p.text:
            p.text = p.text.replace('NOMBRE_CLIENTE', cliente.nombre.upper() + " " + cliente.apellido.upper())
        if 'DOCUMENTO_CLIENTE' in p.text:
            p.text = p.text.replace('DOCUMENTO_CLIENTE', doc_num)
        if 'TELEFONO_CLIENTE' in p.text:
            p.text = p.text.replace('TELEFONO_CLIENTE', telefono)
        if 'DIRECCION_CLIENTE' in p.text:
            p.text = p.text.replace('DIRECCION_CLIENTE', domicilio)
        if 'MOTO_MARCA' in p.text:
            p.text = p.text.replace('MOTO_MARCA', moto.marca)
        if 'MOTO_MODELO' in p.text:
            p.text = p.text.replace('MOTO_MODELO', moto.modelo)
        if 'COLOR_MOTO' in p.text:
            p.text = p.text.replace('COLOR_MOTO', moto.color)
        if 'MOTO_CILINDROS' in p.text:
            p.text = p.text.replace('MOTO_CILINDROS', str(moto.num_cilindros))
        if 'MOTO_MOTOR' in p.text:
            p.text = p.text.replace('MOTO_MOTOR', str(moto.motor))
        if 'MOTO_PASAJEROS' in p.text:
            p.text = p.text.replace('MOTO_PASAJEROS', str(moto.cantidad_pasajeros))
        if 'MOTO_ANIO' in p.text:
            p.text = p.text.replace('MOTO_ANIO', str(moto.anio))
        if 'MOTO_NUM_MOTOR' in p.text:
            p.text = p.text.replace('MOTO_NUM_MOTOR', numero_motor)
        if 'MOTO_NUM_CHASIS' in p.text:
            p.text = p.text.replace('MOTO_NUM_CHASIS', numero_chasis)
        if 'MOTO_MONEDA' in p.text:
            p.text = p.text.replace('MOTO_MONEDA', moneda)
        if 'PRECIO_MOTO' in p.text:
            p.text = p.text.replace('PRECIO_MOTO', str(moto.precio))

    nombre_archivo = f"{moto.id}_{moto.marca}_{moto.modelo}_{cliente.nombre}_{cliente.apellido}_{id_cv}"
    docx_file_path = os.path.join(output_dir, f"{nombre_archivo}.docx")
    doc.save(docx_file_path)

    ruta_archivo = convertir_docx_a_pdf(docx_file_path,nombre_archivo,id_cv)
    if os.path.exists(docx_file_path):
        os.remove(docx_file_path)


def convertir_docx_a_pdf(docx_file_path,nombre_archivo,id_cv):
    # Inicializar COM
    pythoncom.CoInitialize()

    try:
        # Iniciar una instancia de Word
        word = win32com.client.Dispatch("Word.Application")
        
        # Abrir el archivo .docx
        doc = word.Documents.Open(docx_file_path)

        # Asegurarse de que el directorio para guardar el PDF exista
        output_dir = os.path.join(settings.MEDIA_ROOT, 'certificado_venta')
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        # Definir la ruta del archivo PDF de salida
        pdf_file_path = os.path.join(output_dir, f'{nombre_archivo}.pdf')

        # Guardar el documento como .pdf
        doc.SaveAs(pdf_file_path, FileFormat=17)  # El formato 17 corresponde a PDF
        
        # Cerrar el documento y la aplicación de Word
        doc.Close()
        word.Quit()

        # Eliminar el archivo .docx original después de generar el PDF
        if os.path.exists(docx_file_path):
            os.remove(docx_file_path)
            print(f"El archivo .docx '{docx_file_path}' ha sido eliminado.")

        # Devolver la ruta del archivo PDF generado
        compra_venta = ComprasVentas.objects.get(id=id_cv)
    
    # Abrir el archivo PDF y asignarlo al campo certificado_venta
        with open(pdf_file_path, 'rb') as pdf_file:
            compra_venta.certificado_venta.save(f'{nombre_archivo}.pdf', File(pdf_file))
    
    # Guardar los cambios en la base de datos
        compra_venta.save()

        return pdf_file_path

    except Exception as e:
        print(f"Error: {e}")
    finally:
        # Finalizar COM (si es necesario)
        pythoncom.CoUninitialize()

def contexto_editar_usuario(req,mensaje_error):
    usuario = req.user
    usuario_actual = Personal.objects.filter(usuario=usuario.username).first()
    f_nac_formateada = usuario_actual.fecha_nacimiento.strftime('%Y-%m-%d')
    tipo_documento_ci = usuario_actual.documento[0:2:1]
    tipo_documento_pas_dni = usuario_actual.documento[0:3:1]

    longitud_doc = len(usuario_actual.documento)
    doc_num = ""
    
    if tipo_documento_ci == "CI":
        #
        tipo_doc = "CI"
        for i in range(2,longitud_doc):
            doc_num = doc_num + usuario_actual.documento[i]
    elif tipo_documento_pas_dni == "DNI":
        tipo_doc = "DNI"
        for i in range(3,longitud_doc):
            doc_num = doc_num + usuario_actual.documento[i]
    else:
        tipo_doc = "PAS"
        for i in range(3,longitud_doc):
            doc_num = doc_num + usuario_actual.documento[i]

    longitud_correo_princ = len(usuario_actual.correo)
    c_p = "" 
    lon = 0
    for i in range(0,longitud_correo_princ):
        if usuario_actual.correo[i] == "@":
            j = i
            break
        else:
            lon += 1
            c_p = c_p + usuario_actual.correo[i]
    
    longitud_dominio = longitud_correo_princ - lon
    dom_princ = ""
    pos = j
    for j in range(pos,longitud_correo_princ):
        dom_princ = dom_princ + usuario_actual.correo[j]
    
    contexto = {"datos":usuario_actual,
                "f_nac":f_nac_formateada,
                "tipo_doc":tipo_doc,
                "doc_num":doc_num,
                "correo":c_p,
                "dom_princ":dom_princ,
                "error_message":mensaje_error if mensaje_error else None}
    return contexto

def json_para_resumen_pagos(moto,id_cv):
    #SE PONE inicial = 0 PARA QUE NO MUESTRE LA PRIMER "FINANCIACION" QUE SE GUARDA AL EJECUTAR LA VENTA
    financiaciones = Financiamientos.objects.filter(venta_id=id_cv,inicial=0).order_by('fecha')
    cv = ComprasVentas.objects.get(id=id_cv)
    primer_financiacion = Financiamientos.objects.filter(venta_id=id_cv,inicial=1).first()
    dato_inicial = [
            {
                "detalle":moto.marca + " " + moto.modelo,
                "precio_inicial":int(moto.precio) ,
                "precio_final":int(moto.precio_final),
                "moneda":moto.moneda_precio,
                "fecha_compra":cv.fecha_compra.strftime('%d-%m-%Y')
                # "financiacion":f"{primer_financiacion.cantidad_cuotas} x {primer_financiacion.moneda_cuota} {primer_financiacion.valor_cuota}"
            }
        ]
    #
    # if financiaciones:
    #     fin_data = [
    #         {
    #             "detalle": moto.marca + " " + moto.modelo,
    #             "precio_inicial":int(fin.precio_moto_actual),
    #             "precio_final": int(fin.cantidad_cuotas * fin.valor_cuota),
    #             "moneda": fin.moneda_cuota,
    #             "financiacion": f"{fin.cantidad_cuotas} x {fin.moneda_cuota} {fin.valor_cuota}"
    #         }
    #         for fin in financiaciones
    #     ]
    #     union = dato_inicial + fin_data
    #     fin_pagos_json = json.dumps(union)
    
    # union = dato_inicial + fin_data if financiaciones else dato_inicial
    fin_pagos_json = json.dumps(dato_inicial)
    


    primeros_pagos = CuotasMoto.objects.filter(venta_id=id_cv,tipo_pago__in=["Seña", "Entrega inicial","Entrega"]).order_by('-fecha_pago')
    if primeros_pagos:
        p_pagos_data = [
                {   
                    "tipo_pago":cuota.tipo_pago,
                    "fecha":cuota.fecha_pago.strftime('%Y-%m-%d'),
                    "moneda":cuota.moneda,
                    "monto": float(cuota.valor_pago_pesos) if cuota.moneda == "Pesos" else float(cuota.valor_pago_dolares),
                    # "fecha_vencimiento": cuota.fecha_prox_pago.strftime('%Y-%m-%d'),
                }
                for cuota in primeros_pagos
            ]
        primeros_pagos_json = json.dumps(p_pagos_data)
    else:
        primeros_pagos_json = None
    
    cuotas = CuotasMoto.objects.filter(venta_id=id_cv).exclude(tipo_pago__in=["Seña", "Entrega inicial","Entrega"]).order_by('-id')

    if cuotas:
        cuotas_data = [
                {  
                    "tipo_pago":cuota.tipo_pago,
                    "fecha":cuota.fecha_pago.strftime('%Y-%m-%d'),
                    "moneda":cuota.moneda,
                    "monto": float(cuota.valor_pago_pesos) if cuota.moneda == "Pesos" else float(cuota.valor_pago_dolares),
                    "fecha_vencimiento": cuota.fecha_prox_pago.strftime('%Y-%m-%d'),
                }
                for cuota in cuotas
            ]
        cuotas_json = json.dumps(cuotas_data)
    else:
        cuotas_json = None
    
    cliente = Cliente.objects.get(id=cv.cliente_id)
    telefono = ClienteTelefono.objects.filter(cliente_id=cv.cliente_id,principal=1).first()
    # apto = "Apto " + str(cliente.num_apartamento) if int(cliente.num_apartamento) > 0 else ""
    cliente_data = {
        "cliente": cliente.nombre + " " + cliente.apellido,
        "telefono": telefono.telefono,
        "direccion": cliente.domicilio,
        }
    cliente_json = json.dumps(cliente_data)
    logo = Logos.objects.get(id=1)
    financiamiento_actual = Financiamientos.objects.filter(venta_id=id_cv,inicial=0,actual=1).first()
    if financiamiento_actual:
        if financiamiento_actual.moneda_cuota == "Pesos":
            moneda_financiamiento = "$"
        else:
            moneda_financiamiento = "U$S"
    detalle_data={
        "logo_empresa":logo.logo_UM.url,
        "detalle":moto.marca + " " + moto.modelo,
        "precio_inicial":float(moto.precio),
        "precio_final":float(moto.precio_final),
        "financiacion":str(financiamiento_actual.cantidad_cuotas) + " x " + moneda_financiamiento + str(financiamiento_actual.valor_cuota) if financiamiento_actual else "Sin financiamiento"
    }
    detalle_json = json.dumps(detalle_data)

    datos = [
        cuotas_json,
        cliente_json,
        detalle_json,
        primeros_pagos_json,
        fin_pagos_json
    ]

    return datos


def funcion_detalles_cuotas(req,id_cv,buscar,id_buscar_f):
    datos = obtener_detalles_cuotas_comunes(id_cv)
    #
    
    refinanciaciones = Financiamientos.objects.filter(venta_id=id_cv,inicial = 0).first()
    existen_refinanciaciones = True if refinanciaciones else False #SIRVE PARA MOSTRAR LOS DATOS DE LAS REFINANCIACIONES EN CASO DE QUE EXISTAN
    # mostrar_boton_borrar_financiacion = True if refinanciaciones.actual else False
    if buscar:#SI SE ESTA BUSCANDO UNA REFINANCIACION ESPECIFICA
        page_obj = obtener_detalles_cuotas_financiamiento(req,id_buscar_f)
        financiamiento_buscado = Financiamientos.objects.get(id=id_buscar_f)
        id_f = id_buscar_f
        moneda = "$" if financiamiento_buscado.moneda_cuota == "Pesos" else "U$S" if financiamiento_buscado else None
        financiamiento = f"{str(financiamiento_buscado.cantidad_cuotas)} x {moneda} {str(financiamiento_buscado.valor_cuota)}" if financiamiento_buscado else None
        fecha = financiamiento_buscado.fecha if financiamiento_buscado else None
        mensual_quincena = "Quincena" if financiamiento_buscado.quincena else "Mensual"
        mostrar_boton = True if financiamiento_buscado.actual else False
        moneda_pago = financiamiento_buscado.moneda_cuota if financiamiento_buscado else None
        if financiamiento_buscado:
            if financiamiento_buscado.quincena:
                monto_cuota = int(float(financiamiento_buscado.valor_cuota) / 2)
            else:
                monto_cuota = int(financiamiento_buscado.valor_cuota)
        else:
            monto_cuota = None
        mostrar_boton_borrar_financiacion = True if financiamiento_buscado.actual else False
    else:
        fin_actual = Financiamientos.objects.filter(venta_id=id_cv,actual=1).first()
        if fin_actual:
            mostrar_boton = True if fin_actual.actual else False
            page_obj = obtener_detalles_cuotas_financiamiento(req,fin_actual.id)
            moneda = "$" if fin_actual.moneda_cuota == "Pesos" else "U$S" if fin_actual else None
            financiamiento = f"{str(fin_actual.cantidad_cuotas)} x {moneda} {str(fin_actual.valor_cuota)}" if fin_actual else None
            fecha = fin_actual.fecha if fin_actual else None
            id_f = fin_actual.id if fin_actual else None
            moneda_pago = fin_actual.moneda_cuota if fin_actual else None
            if fin_actual:
                if fin_actual.quincena:
                    monto_cuota = int(float(fin_actual.valor_cuota) / 2)
                else:
                    monto_cuota = int(fin_actual.valor_cuota)
            else:
                monto_cuota = None
            mostrar_boton_borrar_financiacion = True if fin_actual.actual else False
            mensual_quincena = "Quincena" if fin_actual.quincena else "Mensual"
        else:
            #SI NO EXISTE NINGUNA REFINANCIACION NO MUESTRA NADA
            # fins = Financiamientos.objects.filter(venta_id=id_cv,inicial=0).first()
            # #BOTON PAGO REFINANCIAMIENTO
            # mostrar_boton = True if fins.actual else False
            # page_obj = obtener_detalles_cuotas_financiamiento(req,fins.id)
            # moneda = "$" if fins.moneda_cuota == "Pesos" else "U$S" if fins else None
            # financiamiento = f"{str(fins.cantidad_cuotas)} x {moneda} {str(fins.valor_cuota)}" if fins else None
            # fecha = fins.fecha if fins else None
            # id_f = fins.id if fins else None
            financiacion_anterior = Financiamientos.objects.filter(venta_id=id_cv,inicial=0).first()
            if financiacion_anterior:
                financiacion_anterior = Financiamientos.objects.filter(venta_id=id_cv,inicial=0).latest('id')
                mostrar_boton = False
                page_obj = obtener_detalles_cuotas_financiamiento(req,financiacion_anterior.id)
                moneda = "$" if financiacion_anterior.moneda_cuota == "Pesos" else "U$S" if financiacion_anterior else None
                financiamiento = f"{str(financiacion_anterior.cantidad_cuotas)} x {moneda} {str(financiacion_anterior.valor_cuota)}" if financiacion_anterior else None
                fecha = financiacion_anterior.fecha if financiacion_anterior else None
                id_f = financiacion_anterior.id if financiacion_anterior else None
                moneda_pago = financiacion_anterior.moneda_cuota if financiacion_anterior else None
                # monto_cuota = int(financiacion_anterior.valor_cuota) if financiacion_anterior else None
                if financiacion_anterior:
                    if financiacion_anterior.quincena:
                        monto_cuota = int(float(financiacion_anterior.valor_cuota) / 2)
                    else:
                        monto_cuota = int(financiacion_anterior.valor_cuota)
                else:
                    monto_cuota = None
                mostrar_boton_borrar_financiacion = False
                mensual_quincena = "Quincena" if financiacion_anterior.quincena else "Mensual"
            else:
                mostrar_boton = False
                page_obj = False
                moneda = False
                financiamiento = False
                fecha = False
                id_f = False
                moneda_pago = None
                monto_cuota = None
                mostrar_boton_borrar_financiacion = False
                mensual_quincena = None
    
    

    fin_inicial = Financiamientos.objects.filter(venta_id=id_cv,inicial=1).first()
    # if fin_actual:
    #     moneda = "$" if fin_actual.moneda_cuota == "Pesos" else "U$S"
    #     # fin_json = f"{str(fin_actual.cantidad_cuotas)} x {moneda} {str(fin_actual.valor_cuota)}"
    # elif fin_inicial:
    #     moneda_ini = "$" if fin_inicial.moneda_cuota == "Pesos" else "U$S"
    #     # fin_json = f"{str(fin_inicial.cantidad_cuotas)} x {moneda_ini} {str(fin_inicial.valor_cuota)}"
    # else:
    #     pass
    #     # fin_json = "Sin financiamiento"
    if fin_inicial:
        moneda_ini = "$" if fin_inicial.moneda_cuota == "Pesos" else "U$S"
        fin_ini = f"{str(fin_inicial.cantidad_cuotas)} x {moneda_ini} {str(fin_inicial.valor_cuota)}"

    ult_cuota = CuotasMoto.objects.filter(venta_id=id_cv).first()
    precio_dolar = PrecioDolar.objects.get(id=1)
    
    cv = ComprasVentas.objects.get(id=id_cv)
    moto = Moto.objects.get(id=cv.moto_id)
    if ult_cuota:
        ult_cuota = CuotasMoto.objects.filter(venta_id=id_cv).latest('id')
        cant_restante_pesos = ult_cuota.cant_restante_pesos
        cant_restante_dolares = ult_cuota.cant_restante_dolares
    else:
        
        if moto.moneda_precio == "Pesos":
            cant_restante_pesos = moto.precio
            cant_restante_dolares = int(moto.precio / precio_dolar.precio_dolar_tienda)
        else:
            cant_restante_pesos = int(moto.precio * precio_dolar.precio_dolar_tienda)
            cant_restante_dolares = moto.precio

    data_jason = json_para_resumen_pagos(moto,id_cv)
    precio = float(precio_dolar.precio_dolar_tienda) if precio_dolar.precio_dolar_tienda else 0
    contexto = {"id_cv":id_cv,
                "producto":datos[0],
                "precio_inicial":datos[1],
                "precio_final":datos[2],
                "pago_acordado":datos[3],
                "primeros_pagos":datos[4],
                "financiamiento":datos[5],
                # "financiamientos":datos[6],
                "page_obj":page_obj,
                "financiacion_info":financiamiento,
                # "financiacion_inicial":fin_ini,
                "fecha_financiacion":fecha,
                "boton_pagar":mostrar_boton,
                "fin_actual":existen_refinanciaciones,
                "cant_restante_dolares":int(cant_restante_dolares),
                "cant_restante_pesos":int(cant_restante_pesos),
                "precio_dolar":precio,
                "id_cliente":cv.cliente_id,
                'cuotas_json': data_jason[0],
                'cliente_json': data_jason[1],
                "detalle_json":data_jason[2],
                "p_pagos_json":data_jason[3],
                "id_f":id_f,
                "mon_pago":moneda_pago,
                "fin_pagos_json":data_jason[4],
                "monto_cuota":monto_cuota,
                "mostrar_boton_borrar_financiacion":mostrar_boton_borrar_financiacion,
                "mensual_quincena":mensual_quincena}
    # datos = [
    #     cuotas_json,
    #     cliente_json,
    #     detalle_json,
    #     primeros_pagos_json,
    #     fin_pagos_json
    # ]
    
    return contexto

def crear_num_motor():
    moto = Moto.objects.filter(num_motor__contains="SNM").count()
    moto = moto + 1
    numero_motor = "SNM" + str(moto)
    return numero_motor

def crear_num_chasis():
    moto = Moto.objects.filter(num_chasis__contains="SNC").count()
    moto = moto + 1
    numero_motor = "SNC" + str(moto)
    return numero_motor

def enviar_correo(titulo,mensaje,destino):
    send_mail(
            subject=titulo,
            message=mensaje,
            #CAMBIAR POR CORREO DE UMPIERREZ MOTOS
            from_email='lumacajuanmanuel@gmail.com',
            #CAMBIAR POR CORREO DE CLIENTE (tiene_correo.correo)
            recipient_list=[destino],
        )

def json_para_servicio(id_s):
    servicio = Servicios.objects.get(id=id_s)
    moto = Moto.objects.get(id=servicio.moto_id)
    cliente = Cliente.objects.get(id=servicio.cliente_id)
    # dato_fijos = [
    #         {
    #             "detalle":moto.marca + " " + moto.modelo,
    #             "cliente":cliente.nombre + " " + cliente.apellido,
    #             "precio":int(servicio.precio),
    #             "fecha":servicio.fecha_ingreso.strftime('%d-%m-%Y'),
    #             "tipo_servicio":servicio.titulo
    #         }
    #     ]
    matricula = Matriculas.objects.filter(moto_id=servicio.moto_id).first()
    matr_actual = matricula.matricula if matricula else "No contiene"
    datos_fijos = {
                "detalle":moto.marca + " " + moto.modelo,
                "num_motor":moto.num_motor if moto.contiene_num_motor else "No contiene",
                "num_chasis":moto.num_chasis if moto.contiene_num_chasis else "No contiene",
                "matricula":matr_actual,
                "cliente":cliente.nombre + " " + cliente.apellido,
                "precio":int(servicio.precio),
                "fecha":servicio.fecha_ingreso.strftime('%d-%m-%Y'),
                "fecha":servicio.fecha_prox_servicio.strftime('%d-%m-%Y') if servicio.fecha_prox_servicio else "No contiene",
                "tipo_servicio":servicio.titulo
            }
    json_datos_fijos = json.dumps(datos_fijos)

    tareas_realizadas = TareasServicios.objects.filter(servicio_id=id_s,realizado=1)
    tareas_servicio = [
                {   
                    "tareas":tarea.tarea,
                }
                for tarea in tareas_realizadas
            ]
    json_tareas = json.dumps(tareas_servicio)


    observaciones = AnotacionesServicio.objects.filter(servicio_id=id_s)
    obs_servicio = [
                {   
                    "observaciones":obs.anotaciones,
                }
                for obs in observaciones
            ]
    json_observaciones = json.dumps(obs_servicio)

    data = [
        json_datos_fijos,
        json_tareas,
        json_observaciones
    ]
    return data

def generar_compromiso_compra_venta(req,id_moto,id_cliente):
    # try:
        cliente = Cliente.objects.get(id=id_cliente)
        moto = Moto.objects.get(id=id_moto)
        telefono_cliente = ClienteTelefono.objects.filter(cliente=cliente,principal=1).first()
        telefono = telefono_cliente.telefono
        numero_letra = num2words(moto.precio, lang='es')
        matricula_padron_moto = Matriculas.objects.filter(moto_id=id_moto).first()
        matricula = matricula_padron_moto.matricula if matricula_padron_moto else ".............................."
        padron = matricula_padron_moto.padron if matricula_padron_moto else ".............................."
        departamento = departamento_matricula(matricula) if matricula_padron_moto else ".............................."

        usuario = req.user
        usuario_actual = Personal.objects.filter(usuario=usuario.username).first()
        nombre_apellido_personal = usuario_actual.nombre + " " + usuario_actual.apellido
        longitud_doc_personal = len(usuario_actual.documento)
        doc_num_personal = ""
        # print("MATRICULA ES: " + matricula)
        for i in range(2,longitud_doc_personal):
            doc_num_personal = doc_num_personal + usuario_actual.documento[i]
        
        if not req.POST['texto_pagaderos']:
            pagaderos = "........................................................................................................................"
        else:
            pagaderos = req.POST['texto_pagaderos']
        # tercero = req.POST['texto_tercero']
        output_dir = os.path.join(settings.MEDIA_ROOT, 'documentacion/compra_venta')

        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

   
        plantilla_path = os.path.join(settings.BASE_DIR, 'media', 'COMPROMISO_COMPRA_VENTA.docx')
        
        # Crear el documento Word a partir del archivo de plantilla
        doc = Document(plantilla_path)

        locale.setlocale(locale.LC_TIME, 'spanish')
        fecha_actual = datetime.now()
        fecha_formateada = fecha_actual.strftime("%d de %B de %Y")
        longitud_doc = len(cliente.documento)
        doc_num = ""
        # print("MATRICULA ES: " + matricula)
        for i in range(2,longitud_doc):
            doc_num = doc_num + cliente.documento[i]
        
        # domicilio = cliente.domicilio.upper()
        if moto.contiene_num_motor:
            numero_motor = moto.num_motor
        else:
            numero_motor = "SIN NUMERO DE MOTOR"

        if moto.contiene_num_chasis:
            numero_chasis = moto.num_chasis
        else:
            numero_chasis = "SIN NUMERO DE CHASIS"
        for p in doc.paragraphs:
            for run in p.runs:
                reemplazado = False  # Variable para rastrear si hubo reemplazo en este run
                if 'personal_nombre_apellido' in run.text:
                    run.text = run.text.replace('personal_nombre_apellido', nombre_apellido_personal)
                    reemplazado = True
                if 'personal_documento' in run.text:
                    run.text = run.text.replace('personal_documento', doc_num_personal)
                    reemplazado = True
                if 'OTRO_NOMBRE' in run.text:
                    run.text = run.text.replace('OTRO_NOMBRE', matricula)
                    reemplazado = True
                if 'MATRICULA_DPTO' in run.text:
                    # print("ENTRA")
                    run.text = run.text.replace('MATRICULA_DPTO', departamento)
                    reemplazado = True
                if 'padron_numero' in run.text:
                    print("ENTRA")
                    run.text = run.text.replace('padron_numero', str(padron))
                    reemplazado = True
                if 'FECHA_ACTUAL' in run.text:
                    run.text = run.text.replace('FECHA_ACTUAL', fecha_formateada)
                    reemplazado = True
                if 'NOMBRE_CLIENTE' in run.text:
                    run.text = run.text.replace('NOMBRE_CLIENTE', cliente.nombre + " " + cliente.apellido)
                    reemplazado = True
                if 'DOCUMENTO_CLIENTE' in run.text:
                    run.text = run.text.replace('DOCUMENTO_CLIENTE', doc_num)
                    reemplazado = True
                if 'DOMICILIO_CLIENTE' in run.text:
                    run.text = run.text.replace('DOMICILIO_CLIENTE', cliente.domicilio)
                    reemplazado = True
                if 'TELEFONO_CLIENTE' in run.text:
                    run.text = run.text.replace('TELEFONO_CLIENTE', telefono)
                    reemplazado = True
                if 'MOTO_TIPO' in run.text:
                    run.text = run.text.replace('MOTO_TIPO', moto.tipo)
                    reemplazado = True
                if 'MOTO_MARCA' in run.text:
                    run.text = run.text.replace('MOTO_MARCA', moto.marca)
                    reemplazado = True
                if 'MOTO_MODELO' in run.text:
                    run.text = run.text.replace('MOTO_MODELO', moto.modelo)
                    reemplazado = True
                if 'MOTO_MOTOR' in run.text:
                    run.text = run.text.replace('MOTO_MOTOR', str(moto.motor))
                    reemplazado = True
                if 'MOTO_ANIO' in run.text:
                    run.text = run.text.replace('MOTO_ANIO', str(moto.anio))
                    reemplazado = True
                if 'MOTO_NUM_MOTOR' in run.text:
                    run.text = run.text.replace('MOTO_NUM_MOTOR', numero_motor)
                    reemplazado = True
                if 'MOTO_NUM_CHASIS' in run.text:
                    run.text = run.text.replace('MOTO_NUM_CHASIS', numero_chasis)
                    reemplazado = True
                if 'PRECIO_MOTO' in run.text:
                    run.text = run.text.replace('PRECIO_MOTO', str(moto.precio))
                    reemplazado = True
                if 'PRECIO_LETRAS' in run.text:
                    run.text = run.text.replace('PRECIO_LETRAS', numero_letra)
                    reemplazado = True
                if 'PAGADEROS' in run.text:
                    run.text = run.text.replace('PAGADEROS', pagaderos)
                    reemplazado = True
                
                # if 'DEPARTAMENTO_SUS' in run.text:
                #     if matricula_padron_moto:
                #         run.text = run.text.replace('DEPARTAMENTO_SUS', departamento)
                #         reemplazado = True
                #     else:
                #         run.text = run.text.replace('DEPARTAMENTO_SUS', "….......................................")
                #         reemplazado = True
                # if 'PADRON_SUS' in run.text:
                #     if matricula_padron_moto:
                #         run.text = run.text.replace('PADRON_SUS', padron)
                #         reemplazado = True
                #     else:
                #         run.text = run.text.replace('PADRON_SUS', "................................")
                #         reemplazado = True


                # Solo aplica formato si hubo reemplazo en este run
                if reemplazado:
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)

                    # Configuración adicional para compatibilidad con Word
                    r = run._element
                    rPr = r.get_or_add_rPr()
                    rFonts = OxmlElement("w:rFonts")
                    rFonts.set(qn("w:ascii"), "Calibri")
                    rFonts.set(qn("w:hAnsi"), "Calibri")
                    rFonts.set(qn("w:eastAsia"), "Calibri")
                    rFonts.set(qn("w:cs"), "Calibri")
                    rPr.append(rFonts)     

        nombre_archivo = f"{moto.id}_{moto.marca}_{moto.modelo}_{cliente.nombre}_{cliente.apellido}_COMPRA_VENTA"
        docx_file_path = os.path.join(output_dir, f"{nombre_archivo}.docx")
        doc.save(docx_file_path)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Configurar la respuesta
        response = HttpResponse(
            buffer,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{moto.id}_{cliente.nombre}_compra_venta.docx"'
        return response

        # ruta_archivo = convertir_docx_a_pdf(docx_file_path,nombre_archivo,id_cv)
        # if os.path.exists(docx_file_path):
        #     os.remove(docx_file_path)
    # except Exception as e:
    #     pass

def contexto_venta_moto(req,id_moto,mensaje_error,documento):
        documento = documento
        cliente = Cliente.objects.filter(documento=documento).first()
        if cliente:
            existe_reserva_moto = ComprasVentas.objects.filter(moto_id=id_moto,tipo="R").first()
            if (existe_reserva_moto) and (existe_reserva_moto.cliente_id != cliente.id):
                contexto = {"datos_moto":False,
                        "error_message":"La moto se encuentra reservada a nombre de otro cliente"}
            else:
                tel1 = ClienteTelefono.objects.filter(principal=1,cliente_id=cliente.id).first()
                tel_1 = tel1.telefono
                tel2 = ClienteTelefono.objects.filter(principal=0,cliente_id=cliente.id).first()

                correo1 = ClienteCorreo.objects.filter(principal=1,cliente_id=cliente.id).first()
                correo2 = ClienteCorreo.objects.filter(principal=0,cliente_id=cliente.id).first()
                if tel2:
                    tel_2 = tel2.telefono
                else:
                    tel_2 = None

                if correo1:
                    c_1 = correo1.correo
                else:
                    c_1 = None
                
                if correo2:
                    c_2 = correo1.correo
                else:
                    c_2 = None
                moto = Moto.objects.get(id=id_moto)

                dolar = PrecioDolar.objects.get(id=1)
                precio_dolar = dolar.precio_dolar_tienda

                existen_fondos = ClienteFondos.objects.filter(cliente=cliente).first()
                if existen_fondos:
                    existen_fondos = ClienteFondos.objects.filter(cliente=cliente).latest('id')
                    if moto.moneda_precio == "Pesos":
                        fondos = int(existen_fondos.total_pesos)
                    else:
                        fondos = int(existen_fondos.total_dolares)
                else:
                    fondos = 0
                
                cv = ComprasVentas.objects.filter(cliente_id=cliente.id,moto_id=id_moto,tipo="R").first()
                if cv:
                    #SI EXISTE UNA RESERVA DE LA MOTO A NOMBRE DEL CLIENTE Y CON PAGOS REALIZADOS
                    pagos = CuotasMoto.objects.filter(venta_id=cv.id)
                    if pagos.exists():
                        suma_total_pesos = 0
                        suma_total_dolares = 0
                        for pago in pagos:
                            suma_total_pesos = suma_total_pesos + int(pago.valor_pago_pesos)
                            suma_total_dolares = suma_total_dolares + int(pago.valor_pago_dolares)
                        
                        if moto.moneda_precio == "Pesos":
                            #SI EL PAGO FUE EN PESOS, EL CAMPO DOLARES ES 0
                            #SI EL PAGO FUE EN DOLARES, EL CAMPO PESOS ES 0
                            suma_total_dolares = float(suma_total_dolares * precio_dolar)
                            total = int(suma_total_dolares + suma_total_pesos)
                        else:
                            suma_total_pesos = float(suma_total_pesos / precio_dolar)
                            total = int(suma_total_dolares + suma_total_pesos)
                    else: 
                        total = 0
                else:
                    total = 0
                
                precio_moto = int(moto.precio) - int(total) - int(fondos)
                if precio_moto < 0:
                    precio_moto = 0
                

                existe_matricula = Matriculas.objects.filter(moto_id=id_moto).first()
                if existe_matricula:
                    matricula = existe_matricula.matricula 
                    padron = existe_matricula.padron
                    departamento = departamento_matricula(matricula)
                else:
                    matricula = None
                    departamento = None
                    padron = None

                #RENDERIZAR PAPEL COMPRA-VENTA
                numero_letra = num2words(moto.precio, lang='es').upper()
                fecha = date.today()
                logo_cv = Logos.objects.get(id=2)
                logo_cv_url = req.build_absolute_uri(logo_cv.logo_UM.url) if logo_cv.logo_UM else None
                precio = int(moto.precio) if moto.precio else 0
                usuario = req.user
                vendedor = Personal.objects.filter(usuario=usuario.username).first()
                nom_ape_vend = vendedor.nombre + " " + vendedor.apellido
                longitud_doc = len(vendedor.documento)
                doc_num = ""
                for i in range(2,longitud_doc):
                    doc_num = doc_num + vendedor.documento[i]
                # print(numero_letra)

                if moto.moneda_precio == "Pesos":
                    precio_en_pesos = int(moto.precio)
                    precio_en_dolares = precio_en_pesos / float(precio_dolar)
                else:
                    precio_en_dolares = int(moto.precio)
                    precio_en_pesos = precio_en_dolares * float(precio_dolar)
                contexto = {"datos_moto":True,
                            "cliente":cliente,
                            "moto":moto,
                            "tel1":tel_1,
                            "tel2":tel_2,
                            "correo1":c_1,
                            "correo2":c_2,
                            "num_letra":numero_letra,
                            "matricula":matricula,
                            "padron":padron,
                            "precio":precio,
                            "departamento":departamento if departamento else None,
                            "fecha":fecha,
                            "logo_cv":logo_cv_url,
                            "vendedor_nombre":nom_ape_vend,
                            "vendedor_ci":doc_num,
                            "fondos":precio_moto,
                            "total_fondos":fondos,
                            "error_message":mensaje_error if mensaje_error else None,
                            "precio_en_dolares":int(precio_en_dolares),
                            "precio_en_pesos":int(precio_en_pesos)}
            # return render(req,"perfil_administrativo/motos/venta_moto.html",{})
        else:
            contexto = {"datos_moto":False,
                        "error_message_cliente":"El cliente no se encuentra registrado en el sistema, para ingresarlo haga clic "}
    
        return contexto

def mi_funcion_diaria():
    clientes = Cliente.objects.all()
    hoy = date.today()  # Obtiene la fecha actual
    dia_hoy = hoy.day
    mes_hoy = hoy.month
    for cliente in clientes:
        if cliente.fecha_nacimiento and (cliente.fecha_nacimiento.day, cliente.fecha_nacimiento.month) == (hoy.day, hoy.month):
            tipo = "Cumpleaños"
            descripcion = f"¡Hoy es el cumpleaños de {cliente.nombre} {cliente.apellido}!"
            insert_notificaciones(descripcion,tipo)
    
    pagos_atrasados = CuotasMoto.objects.all()

    for pagos in pagos_atrasados:
        venta = ComprasVentas.objects.get(id=pagos.venta_id)
        ult_pago = CuotasMoto.objects.filter(venta=venta).latest('id')
        if pagos.fecha_prox_pago and (pagos.fecha_prox_pago.day + 1, pagos.fecha_prox_pago.month) == (hoy.day, hoy.month) and pagos.id == ult_pago.id:
            cliente = Cliente.objects.get(id=venta.cliente_id)
            moto = Moto.objects.get(id=venta.moto_id)
            tipo = "Atraso en cuota"
            descripcion = f"El cliente {cliente.nombre} {cliente.apellido} se encuentra atrasado en el pago de su {moto.marca} {moto.modelo} Codigo {pagos.id}"
            insert_notificaciones(descripcion,tipo)
    
    pagos_atrasados_accesorios = CuotasAccesorios.objects.all()

    for pagos in pagos_atrasados_accesorios:
        venta = ClienteAccesorio.objects.get(id=pagos.venta_id)
        ult_pago = CuotasAccesorios.objects.filter(venta=venta).latest('id')
        if pagos.fecha_prox_pago and (pagos.fecha_prox_pago.day + 1, pagos.fecha_prox_pago.month) == (hoy.day, hoy.month) and pagos.id == ult_pago.id:
            cliente = Cliente.objects.get(id=venta.cliente_id)
            accesorio = Accesorio.objects.get(id=venta.accesorio_id)
            tipo = "Atraso en cuota"
            descripcion = f"El cliente {cliente.nombre} {cliente.apellido} se encuentra atrasado en el pago de su {accesorio.tipo} {accesorio.marca} {accesorio.modelo} Codigo {pagos.id}"
            insert_notificaciones(descripcion,tipo)

def cantidad_solicitudes_no_leidas(req):
    if req.user.is_authenticated:
        usuario = req.user
        usuario_actual = Personal.objects.filter(usuario=usuario.username).first()
        notificaciones = NotificacionPersonal.objects.filter(personal=usuario_actual,leido=0).count()
    else:
        notificaciones = 0
    
    return {'notificaciones': notificaciones}

def contexto_venta_repuesto(req,id_rp,mensaje_error,documento):
        cliente = Cliente.objects.filter(documento=documento).first()
        if cliente:
                tel1 = ClienteTelefono.objects.filter(principal=1,cliente_id=cliente.id).first()
                tel_1 = tel1.telefono
                tel2 = ClienteTelefono.objects.filter(principal=0,cliente_id=cliente.id).first()

                correo1 = ClienteCorreo.objects.filter(principal=1,cliente_id=cliente.id).first()
                correo2 = ClienteCorreo.objects.filter(principal=0,cliente_id=cliente.id).first()
                if tel2:
                    tel_2 = tel2.telefono
                else:
                    tel_2 = None

                if correo1:
                    c_1 = correo1.correo
                else:
                    c_1 = None
                
                if correo2:
                    c_2 = correo1.correo
                else:
                    c_2 = None
                rp = RepuestosPiezas.objects.get(id=id_rp)
                
                # print(numero_letra)
                contexto = {"datos":True,
                            "rp":rp,
                            "cliente":cliente, 
                            "tel1":tel_1,
                            "tel2":tel_2,
                            "correo1":c_1,
                            "correo2":c_2,
                            "error_message":mensaje_error if mensaje_error else None}
            # return render(req,"perfil_administrativo/motos/venta_moto.html",{})
        else:
            contexto = {"datos":False,
                        "error_message_cliente":"El cliente no se encuentra registrado en el sistema, para ingresarlo haga clic "}
    
        return contexto             


def generar_compromiso_compra_venta_moto_ingreso(req,id_moto,id_cliente):
    # try:
        cliente = Cliente.objects.get(id=id_cliente)
        moto = Moto.objects.get(id=id_moto)
        telefono_cliente = ClienteTelefono.objects.filter(cliente=cliente,principal=1).first()
        numero_letra = num2words(moto.precio, lang='es')
        matricula_padron_moto = Matriculas.objects.filter(moto_id=id_moto).first()
        matricula = matricula_padron_moto.matricula if matricula_padron_moto else ".............................."
        padron = matricula_padron_moto.padron if matricula_padron_moto else ".............................."
        departamento = departamento_matricula(matricula) if matricula_padron_moto else ".............................."

        usuario = req.user
        usuario_actual = Personal.objects.filter(usuario=usuario.username).first()
        nombre_apellido_personal = usuario_actual.nombre + " " + usuario_actual.apellido
        # telefono = usuario_actual.telefono
        longitud_doc_personal = len(usuario_actual.documento)
        doc_num_personal = ""
        # print("MATRICULA ES: " + matricula)
        for i in range(2,longitud_doc_personal):
            doc_num_personal = doc_num_personal + usuario_actual.documento[i]
        
        
        # tercero = req.POST['texto_tercero']
        output_dir = os.path.join(settings.MEDIA_ROOT, 'documentacion/compra_venta')

        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

   
        plantilla_path = os.path.join(settings.BASE_DIR, 'media', 'COMPROMISO_COMPRA_VENTA.docx')
        
        # Crear el documento Word a partir del archivo de plantilla
        doc = Document(plantilla_path)

        locale.setlocale(locale.LC_TIME, 'spanish')
        fecha_actual = datetime.now()
        fecha_formateada = fecha_actual.strftime("%d de %B de %Y")
        longitud_doc = len(cliente.documento)
        doc_num = ""
        # print("MATRICULA ES: " + matricula)
        for i in range(2,longitud_doc):
            doc_num = doc_num + cliente.documento[i]
        
        # domicilio = cliente.domicilio.upper()
        if moto.contiene_num_motor:
            numero_motor = moto.num_motor
        else:
            numero_motor = "SIN NUMERO DE MOTOR"

        if moto.contiene_num_chasis:
            numero_chasis = moto.num_chasis
        else:
            numero_chasis = "SIN NUMERO DE CHASIS"
        for p in doc.paragraphs:
            for run in p.runs:
                reemplazado = False  # Variable para rastrear si hubo reemplazo en este run
                if 'personal_nombre_apellido' in run.text:
                    run.text = run.text.replace('personal_nombre_apellido', cliente.nombre + " " + cliente.apellido)
                    reemplazado = True
                    
                if 'personal_documento' in run.text:
                    run.text = run.text.replace('personal_documento', doc_num)
                    reemplazado = True
                    
                if 'OTRO_NOMBRE' in run.text:
                    run.text = run.text.replace('OTRO_NOMBRE', matricula)
                    reemplazado = True
                if 'MATRICULA_DPTO' in run.text:
                    # print("ENTRA")
                    run.text = run.text.replace('MATRICULA_DPTO', departamento)
                    reemplazado = True
                if 'padron_numero' in run.text:
                    print("ENTRA")
                    run.text = run.text.replace('padron_numero', str(padron))
                    reemplazado = True
                if 'FECHA_ACTUAL' in run.text:
                    run.text = run.text.replace('FECHA_ACTUAL', fecha_formateada)
                    reemplazado = True
                if 'NOMBRE_CLIENTE' in run.text:
                    run.text = run.text.replace('NOMBRE_CLIENTE', nombre_apellido_personal)
                    reemplazado = True
                if 'DOCUMENTO_CLIENTE' in run.text:
                    run.text = run.text.replace('DOCUMENTO_CLIENTE', doc_num_personal)
                    reemplazado = True
                if 'DOMICILIO_CLIENTE' in run.text:
                    run.text = run.text.replace('DOMICILIO_CLIENTE', "Reina Martínez 1113")
                    reemplazado = True
                if 'TELEFONO_CLIENTE' in run.text:
                    run.text = run.text.replace('TELEFONO_CLIENTE', "096 117 550")
                    reemplazado = True
                if 'MOTO_TIPO' in run.text:
                    run.text = run.text.replace('MOTO_TIPO', moto.tipo)
                    reemplazado = True
                if 'MOTO_MARCA' in run.text:
                    run.text = run.text.replace('MOTO_MARCA', moto.marca)
                    reemplazado = True
                if 'MOTO_MODELO' in run.text:
                    run.text = run.text.replace('MOTO_MODELO', moto.modelo)
                    reemplazado = True
                if 'MOTO_MOTOR' in run.text:
                    run.text = run.text.replace('MOTO_MOTOR', str(moto.motor))
                    reemplazado = True
                if 'MOTO_ANIO' in run.text:
                    run.text = run.text.replace('MOTO_ANIO', str(moto.anio))
                    reemplazado = True
                if 'MOTO_NUM_MOTOR' in run.text:
                    run.text = run.text.replace('MOTO_NUM_MOTOR', numero_motor)
                    reemplazado = True
                if 'MOTO_NUM_CHASIS' in run.text:
                    run.text = run.text.replace('MOTO_NUM_CHASIS', numero_chasis)
                    reemplazado = True
                if 'PRECIO_MOTO' in run.text:
                    run.text = run.text.replace('PRECIO_MOTO', str(moto.precio))
                    reemplazado = True
                if 'PRECIO_LETRAS' in run.text:
                    run.text = run.text.replace('PRECIO_LETRAS', numero_letra)
                    reemplazado = True
                # if 'PAGADEROS' in run.text:
                #     run.text = run.text.replace('PAGADEROS', pagaderos)
                #     reemplazado = True
                
                # if 'DEPARTAMENTO_SUS' in run.text:
                #     if matricula_padron_moto:
                #         run.text = run.text.replace('DEPARTAMENTO_SUS', departamento)
                #         reemplazado = True
                #     else:
                #         run.text = run.text.replace('DEPARTAMENTO_SUS', "….......................................")
                #         reemplazado = True
                # if 'PADRON_SUS' in run.text:
                #     if matricula_padron_moto:
                #         run.text = run.text.replace('PADRON_SUS', padron)
                #         reemplazado = True
                #     else:
                #         run.text = run.text.replace('PADRON_SUS', "................................")
                #         reemplazado = True


                # Solo aplica formato si hubo reemplazo en este run
                if reemplazado:
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)

                    # Configuración adicional para compatibilidad con Word
                    r = run._element
                    rPr = r.get_or_add_rPr()
                    rFonts = OxmlElement("w:rFonts")
                    rFonts.set(qn("w:ascii"), "Calibri")
                    rFonts.set(qn("w:hAnsi"), "Calibri")
                    rFonts.set(qn("w:eastAsia"), "Calibri")
                    rFonts.set(qn("w:cs"), "Calibri")
                    rPr.append(rFonts)     

        nombre_archivo = f"{moto.id}_{moto.marca}_{moto.modelo}_{cliente.nombre}_{cliente.apellido}_COMPRA_VENTA"
        docx_file_path = os.path.join(output_dir, f"{nombre_archivo}.docx")
        doc.save(docx_file_path)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Configurar la respuesta
        response = HttpResponse(
            buffer,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{moto.id}_{cliente.nombre}_compra_venta.docx"'
        return response

        # ruta_archivo = convertir_docx_a_pdf(docx_file_path,nombre_archivo,id_cv)
        # if os.path.exists(docx_file_path):
        #     os.remove(docx_file_path)
    # except Exception as e:
    #     pass

def contexto_cliente_accesorio(req,mensaje,doc):
        accesorios_seleccionados = req.session["accesorios_json"]

        # Aquí puedes procesar la venta de los accesorios
        # print("Accesorios seleccionados:", accesorios_seleccionados)
        # if mostrar == 1:
        # if req.method == "POST":
        documento = doc
        cliente = Cliente.objects.filter(documento=documento).first()
        if cliente:
            tel1 = ClienteTelefono.objects.filter(principal=1,cliente_id=cliente.id).first()
            tel_1 = tel1.telefono
            tel2 = ClienteTelefono.objects.filter(principal=0,cliente_id=cliente.id).first()

            correo1 = ClienteCorreo.objects.filter(principal=1,cliente_id=cliente.id).first()
            correo2 = ClienteCorreo.objects.filter(principal=0,cliente_id=cliente.id).first()
            if tel2:
                tel_2 = tel2.telefono
            else:
                tel_2 = None

            if correo1:
                c_1 = correo1.correo
            else:
                c_1 = None
            
            if correo2:
                c_2 = correo1.correo
            else:
                c_2 = None
            data_accesorio = []
            print(accesorios_seleccionados)
            accesorios_ids = [int(accs) for accs in accesorios_seleccionados]

# Buscamos todos los accesorios en una sola consulta
            accesorios = Accesorio.objects.filter(id__in=accesorios_ids)
            dolar = PrecioDolar.objects.get(id=1)
            precio_dolar = float(dolar.precio_dolar_tienda)
            precio_total_pesos = 0
            precio_total_dolares = 0
            for accesorio in accesorios:
                  # Convertimos a entero
                a = Accesorio.objects.get(id=accesorio.id)
                print(a.tipo)
                data_accesorio.append({
                    "accesorios":a,
                    "precio":"$" + str(a.precio) if a.moneda_precio == "Pesos" else "U$s" + str(a.precio)
                })
                if a.moneda_precio == "Pesos":
                    precio_total_pesos = precio_total_pesos + float(a.precio)
                    precio_total_dolares = precio_total_pesos / precio_dolar
                else:
                    precio_total_dolares = precio_total_dolares + float(a.precio)
                    precio_total_pesos = precio_total_dolares * precio_dolar
                

            print("LLEGA LINEA 1118")

            fondos = ClienteFondos.objects.filter(cliente_id=cliente.id).first()
            if fondos:
                ult_fondo = ClienteFondos.objects.filter(cliente_id=cliente.id).latest('id')
                total_fondos_pesos = float(ult_fondo.total_pesos)
                total_fondos_dolares = float(ult_fondo.total_dolares)
                precio_restante_pesos = float(precio_total_pesos) - total_fondos_pesos
                precio_restante_dolares = float(precio_total_dolares) - total_fondos_dolares
            else:
                total_fondos_pesos = 0
                total_fondos_dolares = 0
                precio_restante_pesos = 0
                precio_restante_dolares = 0
            
            if (precio_restante_pesos < 0) or (precio_restante_dolares < 0):
                precio_restante_pesos = 0
                precio_restante_dolares = 0
            
            contexto = {
                "datos_accesorio":True,
                "cliente":cliente,
                "accesorios":data_accesorio,
                "tel1":tel_1,
                "tel2":tel_2,
                "correo1":c_1,
                "correo2":c_2,
                "precio_total_pesos":str(round(precio_total_pesos,2)).replace(',', '.'),
                "precio_total_dolares":str(round(precio_total_dolares,2)).replace(',', '.'),
                "precio_restante_pesos":str(round(precio_restante_pesos,2)).replace(',', '.'),
                "precio_restante_dolares":str(round(precio_restante_dolares,2)).replace(',', '.'),
                "cantidad_destinada_pesos":int(total_fondos_pesos),
                "cantidad_destinada_dolares":int(total_fondos_dolares),
                "error_message":mensaje
            }
            
            return contexto
        else:
            contexto = {
                "datos_moto":False,"error_message_cliente":"El cliente no se encuentra registrado en el sistema, para ingresarlo haga clic "
            }
            return contexto

# def sincronizar_bd_servidor_laptop():
#     origen_db = "//192.168.1.100/UmpierrezMotos/db.sqlite3"
#     destino_db = "D:/UmpierrezMotos/UmpierrezMotos/db.sqlite3"
    
#     if not os.path.exists(origen_db):
#         print("❌ No se encontró la base de datos en el servidor. Verifica la conexión.")
#         return  # No sigue con la copia si el archivo no existe

#     try:
#         shutil.copy2(origen_db, destino_db)
#         print("✅ Base de datos sincronizada correctamente.")
#     except Exception as e:
#         print(f"❌ Error al sincronizar la base de datos: {e}")

# def sincronizar_media_servidor_laptop():
#     origen_media = r"\\192.168.1.100\UmpierrezMotos\media"
#     destino_media = r"D:\UmpierrezMotos\UmpierrezMotos\media"

#     for root, _, files in os.walk(origen_media):
#         for file in files:
#             origen_archivo = os.path.join(root, file)
#             destino_archivo = os.path.join(destino_media, file)

#             if not os.path.exists(destino_archivo):  # Solo copiar si no existe en la laptop
#                 shutil.copy2(origen_archivo, destino_archivo)
#                 print(f"✅ Copiado {file} al servidor.")

#     print("✅ Carpeta 'media' sincronizada.")

# def prueba_ejec():
#     print("SE EJECUTA...")

# def ejecutar_cada_5_min():
#     while True:
#         # sincronizar_bd_servidor_laptop()
#         sincronizar_bd_servidor_laptop()
#         time.sleep(300)  # Espera 5 minutos (300 segundos)

# def ejecutar_hilo():
#     hilo_sincronizacion = threading.Thread(target=ejecutar_cada_5_min, daemon=True)
#     hilo_sincronizacion.start()

# # Crear un hilo en segundo plano
# hilo = threading.Thread(target=ejecutar_cada_5_min, daemon=True)
# hilo.start()

def validar_caja_abierta():
    try:
        caja = Caja.objects.filter(estado="Abierto").first()
        ult_caja = Caja.objects.latest('id')
        hoy = date.today()
        if not caja:
            return False
        elif ult_caja.apertura.date() != hoy:
            return False
        else:
            return True
    except Exception as e:
        pass